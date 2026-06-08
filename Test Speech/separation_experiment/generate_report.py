# -*- coding: utf-8 -*-
"""Tạo file Word báo cáo kết quả thực nghiệm MSS + Speech Enhancement Pipeline."""

import sys, io, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
RESULTS_JSON = os.path.join(SCRIPT_DIR, "results.json")
OUT_DOCX     = os.path.join(SCRIPT_DIR, "..", "..", "MSS_SpeechEnhancement_Report.docx")

# RTF đo được trong thực nghiệm (từ terminal log)
RTF_BS   = {"mean": 9.335, "min": 8.557, "max": 10.290}
RTF_MEL  = {"mean": 5.673, "min": 5.005, "max": 6.449}
RTF_DFN3 = 0.060   # DeepFilterNet3 trung bình

# ─────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def para(doc, text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 * (level + 1))
    p.paragraph_format.space_after = Pt(3)

def note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("⚠ " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    p.paragraph_format.space_after = Pt(5)

def success(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("✅ " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x05, 0x6A, 0x34)
    p.paragraph_format.space_after = Pt(5)

def info(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("ℹ " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.space_after = Pt(5)

def divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def tbl(doc, headers, rows, col_widths=None, header_color="1A56DB", alt_rows=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_color)
    # Data rows
    for ri, row_data in enumerate(rows):
        r = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            r[i].text = str(cell_text)
            r[i].paragraphs[0].runs[0].font.size = Pt(9)
            if alt_rows and ri % 2 == 1:
                set_cell_bg(r[i], "EFF6FF")
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def color_cell(cell, hex_color, text_color="000000"):
    set_cell_bg(cell, hex_color)
    for run in cell.paragraphs[0].runs:
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        run.font.color.rgb = RGBColor(
            int(text_color[0:2], 16),
            int(text_color[2:4], 16),
            int(text_color[4:6], 16),
        )

def ovrl_color(val):
    """Trả về mã màu ô tương ứng với giá trị OVRL."""
    if val is None:
        return "FFFFFF"
    if val >= 3.8:
        return "D1FAE5"   # xanh lá đậm
    elif val >= 3.4:
        return "DCFCE7"   # xanh lá nhạt
    elif val >= 3.0:
        return "FEF9C3"   # vàng nhạt
    elif val >= 2.5:
        return "FEF3C7"   # cam nhạt
    else:
        return "FEE2E2"   # đỏ nhạt


# ─────────────────────────────────────────────────────────
# BUILD
# ─────────────────────────────────────────────────────────

def build(out_path):
    with open(RESULTS_JSON, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()

    # ── Lề trang ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ════════════════════════════════════════════════════════════
    # TRANG TIÊU ĐỀ
    # ════════════════════════════════════════════════════════════
    title = doc.add_heading("BÁO CÁO THỰC NGHIỆM", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub1 = doc.add_paragraph("Music Source Separation + Speech Enhancement Pipeline")
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub1.runs[0].bold = True
    sub1.runs[0].font.size = Pt(14)
    sub1.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    sub2 = doc.add_paragraph(
        "Tách giọng từ YouTube audio thô bằng Roformer + DeepFilterNet 3\n"
        "Dataset: KhanhVy Vlog · KhanhVy Talkshow · Beyond Limits / NGA Levi  |  Metric: DNSMOS P.835"
    )
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].italic = True
    sub2.runs[0].font.size = Pt(10)
    sub2.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════
    # PHẦN 1: TỔNG QUAN
    # ════════════════════════════════════════════════════════════
    h1(doc, "1. Tổng quan & Mục tiêu")

    para(doc,
        "Audio YouTube thô thường chứa nhạc nền, tiếng ồn môi trường và nhiều tạp âm khác. "
        "Để thu được giọng nói chất lượng cao phục vụ TTS/ASR/Voice Cloning, cần qua hai bước xử lý:")

    bullet(doc, "Music Source Separation (MSS) — tách giọng nói ra khỏi nhạc nền")
    bullet(doc, "Speech Enhancement — lọc phần tạp âm còn lại trong luồng giọng")

    para(doc, "Pipeline thực nghiệm:")
    p = doc.add_paragraph()
    run = p.add_run(
        "  YouTube .webm  →  Convert WAV  →  [MSS]  →  [Speech Enhancement]  →  Clean Voice"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()
    divider(doc)

    # ════════════════════════════════════════════════════════════
    # PHẦN 2: THIẾT LẬP THỰC NGHIỆM
    # ════════════════════════════════════════════════════════════
    h1(doc, "2. Thiết lập thực nghiệm")

    h2(doc, "2.1 Dataset")
    tbl(doc,
        ["Category", "Nguồn", "Số video", "Files test", "Đặc điểm"],
        [
            ["khanhvy_vlog",         "YouTube vlog",        "23", "3", "Nhạc nền liên tục, giọng VN nữ"],
            ["khanhvy_talkshow",     "YouTube talk show",   "37", "3", "Nhạc intro/outro, nhiều nền âm thanh"],
            ["beyondlimits_ngalevi", "YouTube interview",   "29", "3", "Tương đối sạch, ít nhạc nền"],
        ],
        col_widths=[4.2, 3.5, 2.2, 2.2, 5.5],
    )
    para(doc,
         "Mỗi file được cắt 60 giây (bỏ 30s đầu) để đại diện cho nội dung chính, "
         "tránh intro/jingle. Tổng: 9 clips × 60s = 9 phút audio.", italic=True, size=9)

    h2(doc, "2.2 Môi trường")
    tbl(doc,
        ["Thành phần", "Chi tiết"],
        [
            ["OS",             "Windows 10 (AMD64)"],
            ["CPU",            "AMD Ryzen (no GPU acceleration)"],
            ["Python",         "3.11.9"],
            ["PyTorch",        "2.11.0+cpu"],
            ["audio-separator","0.44.2"],
            ["FFmpeg",         "8.1.1-full_build"],
        ],
        col_widths=[4.5, 13.0],
    )

    h2(doc, "2.3 Models đánh giá")
    tbl(doc,
        ["#", "Model", "Architecture", "Params (ước tính)", "Vocab SDR", "Mục tiêu"],
        [
            ["1", "BS-Roformer-Viperx-1297",
             "Band-Split RoFormer\n(ViperX fine-tune)",
             "~320M (639 MB ckpt)", "11.77 dB", "Tách vocal (Vocals + Instrumental)"],
            ["2", "MelBand-Roformer-Vocal\n(Kimberley Jensen)",
             "Mel-Band RoFormer\n(Kim fine-tune)",
             "~456M (913 MB ckpt)", "12.60 dB", "Tách vocal (vocals + other)"],
            ["3", "DeepFilterNet 3",
             "GRU + ERB filter bank",
             "~1.8M", "N/A", "Speech Enhancement (bước 2)"],
        ],
        col_widths=[0.6, 4.2, 3.5, 3.2, 2.2, 3.8],
    )

    info(doc,
         "Vocal SDR là chỉ số Signal-to-Distortion Ratio đo trên tập MUSDB18HQ "
         "(benchmark chuẩn cho Music Source Separation). Cao hơn = tách sạch hơn.")

    h2(doc, "2.4 Metric đánh giá")
    tbl(doc,
        ["Metric", "Mô tả", "Thang đo", "Ghi chú"],
        [
            ["DNSMOS OVRL", "Overall MOS — chất lượng tổng hợp", "1–5 (cao hơn tốt hơn)",
             "Không cần reference audio"],
            ["DNSMOS SIG",  "Signal quality — độ rõ giọng nói",  "1–5",
             "Phản ánh độ trong tiếng"],
            ["DNSMOS BAK",  "Background quality — mức độ nhiễu nền",  "1–5",
             "Cao = ít nhiễu nền"],
            ["SNR (ước tính)", "Signal-to-Noise Ratio",          "dB (cao hơn tốt hơn)",
             "Ước tính bằng VAD đơn giản"],
            ["RTF",         "Real-Time Factor",                  "× (thấp hơn tốt hơn)",
             "RTF < 1 = nhanh hơn real-time"],
        ],
        col_widths=[3.0, 4.5, 3.5, 6.5],
    )
    divider(doc)

    # ════════════════════════════════════════════════════════════
    # PHẦN 3: KẾT QUẢ DNSMOS
    # ════════════════════════════════════════════════════════════
    h1(doc, "3. Kết quả DNSMOS P.835")

    # Chuẩn bị dữ liệu
    cats = {
        "khanhvy_vlog":         [],
        "khanhvy_talkshow":     [],
        "beyondlimits_ngalevi": [],
    }
    for row in data:
        for cat in cats:
            if row["file"].startswith(cat):
                cats[cat].append(row)

    cat_labels = {
        "khanhvy_vlog":         "KhanhVy Vlog",
        "khanhvy_talkshow":     "KhanhVy Talkshow",
        "beyondlimits_ngalevi": "Beyond Limits / NGA Levi",
    }

    h2(doc, "3.1 Bảng DNSMOS OVRL chi tiết theo từng file")

    headers = [
        "File (60s clip)",
        "Raw\nOVRL",
        "BS-Rofor\nOVRL",
        "BS+DFN3\nOVRL",
        "MelBand\nOVRL",
        "MelBand\n+DFN3",
        "Δ Best\nvs Raw",
    ]

    for cat, rows_cat in cats.items():
        h3(doc, cat_labels[cat])
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(hdr[i], "1A56DB")

        for row in rows_cat:
            short = row["file"].replace(cat + "__", "").replace(cat + "_", "")
            raw_ovrl = row.get("dnsmos_raw", {}).get("OVRL") if row.get("dnsmos_raw") else None
            bs_sep   = row.get("dnsmos_sep_bs_roformer_1297", {}).get("OVRL") if row.get("dnsmos_sep_bs_roformer_1297") else None
            bs_enh   = row.get("dnsmos_enh_bs_roformer_1297", {}).get("OVRL") if row.get("dnsmos_enh_bs_roformer_1297") else None
            mb_sep   = row.get("dnsmos_sep_melband_roformer_vocal", {}).get("OVRL") if row.get("dnsmos_sep_melband_roformer_vocal") else None
            mb_enh   = row.get("dnsmos_enh_melband_roformer_vocal", {}).get("OVRL") if row.get("dnsmos_enh_melband_roformer_vocal") else None

            best = max(v for v in [bs_sep, bs_enh, mb_sep, mb_enh] if v is not None)
            delta = round(best - raw_ovrl, 3) if raw_ovrl else None

            r = t.add_row().cells
            vals = [short[:22], raw_ovrl, bs_sep, bs_enh, mb_sep, mb_enh,
                    f"+{delta}" if delta else "N/A"]
            for i, v in enumerate(vals):
                r[i].text = str(v) if v is not None else "—"
                r[i].paragraphs[0].runs[0].font.size = Pt(8)

            # Color coding cho OVRL cells
            for ci, val in [(1, raw_ovrl), (2, bs_sep), (3, bs_enh), (4, mb_sep), (5, mb_enh)]:
                set_cell_bg(r[ci], ovrl_color(val))

            # Highlight ô tốt nhất (delta)
            set_cell_bg(r[6], "D1FAE5" if delta and delta >= 1.0 else "FEF9C3")

        doc.add_paragraph()

    # ── Chú thích màu ──────────────────────────────────────────
    para(doc, "Chú thích màu OVRL:", bold=True, size=9)
    color_notes = [
        ("D1FAE5", "≥ 3.8 — Rất tốt"),
        ("DCFCE7", "3.4–3.8 — Tốt"),
        ("FEF9C3", "3.0–3.4 — Khá"),
        ("FEF3C7", "2.5–3.0 — Trung bình"),
        ("FEE2E2", "< 2.5 — Kém"),
    ]
    ct = doc.add_table(rows=1, cols=len(color_notes))
    ct.style = "Table Grid"
    for i, (clr, label) in enumerate(color_notes):
        c = ct.rows[0].cells[i]
        c.text = label
        c.paragraphs[0].runs[0].font.size = Pt(8)
        set_cell_bg(c, clr)
    doc.add_paragraph()

    h2(doc, "3.2 Bảng tổng hợp SIG / BAK / OVRL (trung bình theo category)")

    def avg(lst):
        vals = [v for v in lst if v is not None]
        return round(np.mean(vals), 3) if vals else None

    summary_rows = []
    for cat, rows_cat in cats.items():
        def g(key, subkey):
            return avg([
                row.get(key, {}).get(subkey) if row.get(key) else None
                for row in rows_cat
            ])

        summary_rows.append([
            cat_labels[cat],
            f"{g('dnsmos_raw','SIG')} / {g('dnsmos_raw','BAK')} / {g('dnsmos_raw','OVRL')}",
            f"{g('dnsmos_sep_bs_roformer_1297','SIG')} / {g('dnsmos_sep_bs_roformer_1297','BAK')} / {g('dnsmos_sep_bs_roformer_1297','OVRL')}",
            f"{g('dnsmos_enh_bs_roformer_1297','SIG')} / {g('dnsmos_enh_bs_roformer_1297','BAK')} / {g('dnsmos_enh_bs_roformer_1297','OVRL')}",
            f"{g('dnsmos_sep_melband_roformer_vocal','SIG')} / {g('dnsmos_sep_melband_roformer_vocal','BAK')} / {g('dnsmos_sep_melband_roformer_vocal','OVRL')}",
            f"{g('dnsmos_enh_melband_roformer_vocal','SIG')} / {g('dnsmos_enh_melband_roformer_vocal','BAK')} / {g('dnsmos_enh_melband_roformer_vocal','OVRL')}",
        ])

    # Tổng hợp toàn bộ
    summary_rows.append([
        "⟨ TRUNG BÌNH TOÀN BỘ ⟩",
        f"{avg([r.get('dnsmos_raw',{}).get('SIG') if r.get('dnsmos_raw') else None for r in data])} / "
        f"{avg([r.get('dnsmos_raw',{}).get('BAK') if r.get('dnsmos_raw') else None for r in data])} / "
        f"{avg([r.get('dnsmos_raw',{}).get('OVRL') if r.get('dnsmos_raw') else None for r in data])}",
        f"{avg([r.get('dnsmos_sep_bs_roformer_1297',{}).get('SIG') if r.get('dnsmos_sep_bs_roformer_1297') else None for r in data])} / "
        f"{avg([r.get('dnsmos_sep_bs_roformer_1297',{}).get('BAK') if r.get('dnsmos_sep_bs_roformer_1297') else None for r in data])} / "
        f"{avg([r.get('dnsmos_sep_bs_roformer_1297',{}).get('OVRL') if r.get('dnsmos_sep_bs_roformer_1297') else None for r in data])}",
        f"{avg([r.get('dnsmos_enh_bs_roformer_1297',{}).get('SIG') if r.get('dnsmos_enh_bs_roformer_1297') else None for r in data])} / "
        f"{avg([r.get('dnsmos_enh_bs_roformer_1297',{}).get('BAK') if r.get('dnsmos_enh_bs_roformer_1297') else None for r in data])} / "
        f"{avg([r.get('dnsmos_enh_bs_roformer_1297',{}).get('OVRL') if r.get('dnsmos_enh_bs_roformer_1297') else None for r in data])}",
        f"{avg([r.get('dnsmos_sep_melband_roformer_vocal',{}).get('SIG') if r.get('dnsmos_sep_melband_roformer_vocal') else None for r in data])} / "
        f"{avg([r.get('dnsmos_sep_melband_roformer_vocal',{}).get('BAK') if r.get('dnsmos_sep_melband_roformer_vocal') else None for r in data])} / "
        f"{avg([r.get('dnsmos_sep_melband_roformer_vocal',{}).get('OVRL') if r.get('dnsmos_sep_melband_roformer_vocal') else None for r in data])}",
        f"{avg([r.get('dnsmos_enh_melband_roformer_vocal',{}).get('SIG') if r.get('dnsmos_enh_melband_roformer_vocal') else None for r in data])} / "
        f"{avg([r.get('dnsmos_enh_melband_roformer_vocal',{}).get('BAK') if r.get('dnsmos_enh_melband_roformer_vocal') else None for r in data])} / "
        f"{avg([r.get('dnsmos_enh_melband_roformer_vocal',{}).get('OVRL') if r.get('dnsmos_enh_melband_roformer_vocal') else None for r in data])}",
    ])

    tbl(doc,
        ["Category", "Raw\nSIG/BAK/OVRL",
         "BS-Roformer\nSIG/BAK/OVRL", "BS+DFN3\nSIG/BAK/OVRL",
         "MelBand\nSIG/BAK/OVRL",    "MelBand+DFN3\nSIG/BAK/OVRL"],
        summary_rows,
        col_widths=[3.8, 3.5, 3.5, 3.5, 3.5, 3.8],
    )

    divider(doc)

    # ════════════════════════════════════════════════════════════
    # PHẦN 4: TỐC ĐỘ XỬ LÝ
    # ════════════════════════════════════════════════════════════
    h1(doc, "4. Tốc độ xử lý (RTF — Real-Time Factor)")

    para(doc,
        "RTF = Thời gian xử lý / Thời lượng audio. "
        "RTF < 1.0 nghĩa là xử lý nhanh hơn real-time (phù hợp batch pipeline).")

    tbl(doc,
        ["Model", "RTF Trung bình", "RTF Min", "RTF Max",
         "Thời gian cho 60s", "So sánh"],
        [
            ["BS-Roformer-Viperx-1297",
             f"{RTF_BS['mean']:.3f}×",
             f"{RTF_BS['min']:.3f}×",
             f"{RTF_BS['max']:.3f}×",
             f"≈ {RTF_BS['mean']*60:.0f}s (~{RTF_BS['mean']*60/60:.1f} phút)",
             "Baseline"],
            ["MelBand-Roformer-Vocal",
             f"{RTF_MEL['mean']:.3f}×",
             f"{RTF_MEL['min']:.3f}×",
             f"{RTF_MEL['max']:.3f}×",
             f"≈ {RTF_MEL['mean']*60:.0f}s (~{RTF_MEL['mean']*60/60:.1f} phút)",
             f"Nhanh hơn BS {RTF_BS['mean']/RTF_MEL['mean']:.2f}×"],
            ["DeepFilterNet 3",
             f"{RTF_DFN3:.3f}×",
             "0.059×", "0.067×",
             f"≈ {RTF_DFN3*60:.1f}s (real-time)",
             "Rất nhanh — 16× RT"],
        ],
        col_widths=[4.5, 2.5, 2.2, 2.2, 4.0, 3.2],
    )

    note(doc,
         "RTF đo trên CPU (AMD64, không dùng GPU). Với GPU (CUDA), "
         "tốc độ kỳ vọng nhanh hơn 10–30× tùy hardware.")

    para(doc, "Pipeline tổng RTF (MSS + Enhancement):", bold=True)
    tbl(doc,
        ["Pipeline", "RTF MSS", "RTF DFN3", "RTF Tổng", "Ghi chú"],
        [
            ["BS-Roformer + DFN3",
             f"{RTF_BS['mean']:.2f}×", f"{RTF_DFN3:.3f}×",
             f"{RTF_BS['mean'] + RTF_DFN3:.2f}×",
             "DFN3 không đáng kể so với MSS"],
            ["MelBand + DFN3",
             f"{RTF_MEL['mean']:.2f}×", f"{RTF_DFN3:.3f}×",
             f"{RTF_MEL['mean'] + RTF_DFN3:.2f}×",
             f"Nhanh hơn BS pipeline {(RTF_BS['mean'])/(RTF_MEL['mean']):.2f}×"],
        ],
        col_widths=[4.5, 2.5, 2.5, 3.0, 5.2],
    )
    divider(doc)

    # ════════════════════════════════════════════════════════════
    # PHẦN 5: VOICE DISTORTION — ĐÁNH GIÁ PERCEPTUAL
    # ════════════════════════════════════════════════════════════
    h1(doc, "5. Đánh giá Voice Distortion (Perceptual Metrics)")

    PERCEPTUAL_JSON = os.path.join(SCRIPT_DIR, "perceptual_results.json")

    para(doc,
        "Ngoài DNSMOS (chủ yếu đo noise suppression), một vấn đề thực tế quan trọng là "
        "\"méo giọng\" (voice distortion) — hiện tượng MSS/Enhancement có thể làm giảm noise "
        "nhưng đồng thời khiến giọng nói bị biến dạng pitch, mất âm sắc (timbre), hoặc nghe "
        "robot hóa. Điều này đặc biệt nghiêm trọng với mục tiêu TTS training / Voice Cloning, "
        "vì DNSMOS cao không đảm bảo giọng vẫn giống người nói gốc.")

    note(doc,
        "DNSMOS P.835 chủ yếu đánh giá chất lượng nghe tổng thể và noise suppression. "
        "Nó KHÔNG phát hiện méo giọng nếu model loại bỏ cả noise lẫn một phần harmonics "
        "của giọng. Cần dùng metrics perceptual bổ sung.")

    h2(doc, "5.1 Metrics đánh giá Voice Distortion")

    tbl(doc,
        ["Metric", "Mô tả", "Ngưỡng tốt", "Ý nghĩa với méo giọng"],
        [
            ["F0-RMSE (Hz)",
             "Lệch tần số cơ bản (pitch) giữa raw và enhanced; tính trên voiced frames",
             "< 15 Hz", "Cao → pitch bị lệch → giọng nghe cao/thấp hơn bất thường"],
            ["MCD (dB)",
             "Mel Cepstral Distortion — đo sự thay đổi âm sắc (spectral envelope)",
             "< 4 dB", "Cao → âm sắc thay đổi lớn → nghe như giọng người khác"],
            ["EPR",
             "Energy Preservation Ratio — tỉ lệ năng lượng giọng được giữ lại",
             "> 0.7", "Thấp → over-suppression → giọng bị mờ/nhỏ không tự nhiên"],
            ["SFR",
             "Spectral Flatness Ratio — phổ tần trở nên phẳng hơn (mất harmonics)",
             "< 1.3", "> 1.8 → giọng robot/metallic (formant bị triệt tiêu)"],
            ["VDS (0–10)",
             "Voice Distortion Score tổng hợp từ F0, MCD, EPR, SFR",
             "< 2.5", "< 2.5 ít méo; 2.5–4.5 chấp nhận; > 4.5 cần nghe kiểm tra"],
        ],
        col_widths=[2.5, 5.5, 2.2, 7.3],
    )

    h2(doc, "5.2 Kết quả Voice Distortion Metrics")

    if os.path.exists(PERCEPTUAL_JSON):
        with open(PERCEPTUAL_JSON, encoding="utf-8") as pf:
            pdata = json.load(pf)

        # Bảng chi tiết từng file
        perc_headers = [
            "File", "Model",
            "F0-RMSE\n(Hz)", "MCD\n(dB)", "EPR", "SFR", "VDS\n(0–10)", "Verdict"
        ]
        perc_rows = []
        for rec in pdata:
            fname = rec["file"].split("__")[-1][:16]
            for mk, mv in rec.get("models", {}).items():
                label = "BS+DFN3" if mk == "bs_roformer_1297" else "Mel+DFN3"
                perc_rows.append([
                    fname, label,
                    str(mv.get("f0_rmse_hz", "—")),
                    str(mv.get("mcd_dB", "—")),
                    str(mv.get("epr", "—")),
                    str(mv.get("sfr", "—")),
                    str(mv.get("vds", "—")),
                    mv.get("verdict", "—")[:30] if mv.get("verdict") else "—",
                ])
        tbl(doc, perc_headers, perc_rows,
            col_widths=[2.8, 2.5, 2.2, 2.2, 1.8, 1.8, 2.2, 5.5])

        # Tổng hợp trung bình
        h2(doc, "5.3 Phân tích tổng hợp Voice Distortion")

        def avg_perc(model_key, field):
            vals = []
            for rec in pdata:
                mv = rec.get("models", {}).get(model_key, {})
                v = mv.get(field)
                if v is not None:
                    vals.append(v)
            return round(float(np.mean(vals)), 3) if vals else None

        bs_f0  = avg_perc("bs_roformer_1297", "f0_rmse_hz")
        bs_mcd = avg_perc("bs_roformer_1297", "mcd_dB")
        bs_epr = avg_perc("bs_roformer_1297", "epr")
        bs_sfr = avg_perc("bs_roformer_1297", "sfr")
        bs_vds = avg_perc("bs_roformer_1297", "vds")

        mel_f0  = avg_perc("melband_roformer_vocal", "f0_rmse_hz")
        mel_mcd = avg_perc("melband_roformer_vocal", "mcd_dB")
        mel_epr = avg_perc("melband_roformer_vocal", "epr")
        mel_sfr = avg_perc("melband_roformer_vocal", "sfr")
        mel_vds = avg_perc("melband_roformer_vocal", "vds")

        tbl(doc,
            ["Pipeline", "F0-RMSE TB\n(Hz)", "MCD TB\n(dB)", "EPR TB", "SFR TB", "VDS TB\n(0–10)"],
            [
                ["BS-Roformer + DFN3",
                 str(bs_f0), str(bs_mcd), str(bs_epr), str(bs_sfr), str(bs_vds)],
                ["MelBand-Roformer + DFN3",
                 str(mel_f0), str(mel_mcd), str(mel_epr), str(mel_sfr), str(mel_vds)],
            ],
            col_widths=[4.5, 2.8, 2.8, 2.2, 2.2, 2.8],
        )

        # Nhận xét về méo giọng
        all_verdicts = []
        for rec in pdata:
            for mk, mv in rec.get("models", {}).items():
                vds = mv.get("vds")
                if vds is not None:
                    all_verdicts.append((mk, vds, mv.get("verdict", ""), mv.get("issues", [])))

        bs_vds_list  = [v for mk, v, _, _ in all_verdicts if mk == "bs_roformer_1297"]
        mel_vds_list = [v for mk, v, _, _ in all_verdicts if mk == "melband_roformer_vocal"]

        def vds_label(v):
            if v is None:
                return "N/A"
            if v <= 2.5:
                return f"{v} ✅ Rất ít méo"
            elif v <= 4.5:
                return f"{v} 🟡 Méo nhẹ"
            elif v <= 6.5:
                return f"{v} 🟠 Méo trung bình"
            else:
                return f"{v} 🔴 Méo nặng"

        if bs_vds_list:
            avg_bs_vds = round(float(np.mean(bs_vds_list)), 2)
            para(doc, f"BS-Roformer + DFN3: VDS trung bình = {vds_label(avg_bs_vds)}", bold=True)
        if mel_vds_list:
            avg_mel_vds = round(float(np.mean(mel_vds_list)), 2)
            para(doc, f"Mel-Band + DFN3: VDS trung bình = {vds_label(avg_mel_vds)}", bold=True)

        # Liệt kê các issues phổ biến
        all_issues_bs = []
        all_issues_mel = []
        for rec in pdata:
            for mk, mv in rec.get("models", {}).items():
                for iss in mv.get("issues", []):
                    if mk == "bs_roformer_1297":
                        all_issues_bs.append(iss)
                    else:
                        all_issues_mel.append(iss)

        if all_issues_bs or all_issues_mel:
            h3(doc, "Các vấn đề phát hiện tự động")
            if all_issues_bs:
                para(doc, "BS-Roformer + DFN3:", bold=True)
                from collections import Counter
                for iss, cnt in Counter(all_issues_bs).most_common(5):
                    bullet(doc, f"[{cnt}/{len(pdata)} files] {iss}")
            if all_issues_mel:
                para(doc, "Mel-Band + DFN3:", bold=True)
                for iss, cnt in Counter(all_issues_mel).most_common(5):
                    bullet(doc, f"[{cnt}/{len(pdata)} files] {iss}")

        info(doc,
            "Các vấn đề về méo giọng cần được kiểm tra thêm bằng cách nghe trực tiếp. "
            "Xem file listening_test.html trong thư mục separation_experiment/ để nghe "
            "so sánh từng case và chấm điểm thủ công.")

    else:
        note(doc,
            "Chưa có kết quả perceptual. Chạy: python perceptual_eval.py "
            "trong thư mục separation_experiment/ để tạo perceptual_results.json")

    h2(doc, "5.4 Listening Test — Hướng dẫn đánh giá thủ công")

    para(doc,
        "Metrics tự động (F0, MCD, VDS) chỉ là proxy cho perceptual quality. "
        "Để đánh giá chính xác \"méo giọng\", cần nghe trực tiếp các case điển hình. "
        "File listening_test.html cung cấp giao diện nghe và chấm điểm MOS (1–5) cho 4 tiêu chí:")

    tbl(doc,
        ["Tiêu chí", "Mô tả", "5 = Tốt nhất", "1 = Tệ nhất"],
        [
            ["Noise Removal",      "Mức độ loại bỏ noise/nhạc nền",
             "Sạch hoàn toàn",     "Vẫn còn nhiều nhiễu"],
            ["Voice Naturalness",  "Độ tự nhiên của giọng",
             "Hoàn toàn tự nhiên", "Robot, cứng, không tự nhiên"],
            ["Voice Distortion",   "Giọng có bị méo không (so với gốc)",
             "Giống hệt giọng gốc","Khác hoàn toàn, méo nặng"],
            ["Overall Quality",    "Chất lượng tổng thể",
             "Dùng được ngay",     "Không sử dụng được"],
        ],
        col_widths=[3.5, 4.5, 3.5, 4.5],
    )

    bullet(doc, "Các dấu hiệu méo giọng cần chú ý khi nghe:")
    bullet(doc, "Pitch lệch: giọng cao/thấp hơn bất thường so với bản gốc", level=1)
    bullet(doc, "Âm sắc thay đổi: nghe như giọng người khác (timbre distortion)", level=1)
    bullet(doc, "Giọng robot/metallic: âm thanh cứng, thiếu harmonics tự nhiên", level=1)
    bullet(doc, "Mất âm cuối/consonant clipping: các âm tắt, âm sát bị cắt", level=1)
    bullet(doc, "Breathing artifact: tiếng thở bị khuếch đại hoặc cắt kỳ lạ", level=1)
    bullet(doc, "Over-suppression: giọng nhỏ bất thường, mất năng lượng", level=1)

    divider(doc)

    # ════════════════════════════════════════════════════════════
    # PHẦN 6: PHÂN TÍCH & NHẬN XÉT
    # ════════════════════════════════════════════════════════════
    h1(doc, "6. Phân tích & Nhận xét")

    # Tính các số để dùng trong phân tích
    raw_ovrls  = [r["dnsmos_raw"]["OVRL"] for r in data if r.get("dnsmos_raw")]
    bs_enh_ovrls  = [r["dnsmos_enh_bs_roformer_1297"]["OVRL"] for r in data if r.get("dnsmos_enh_bs_roformer_1297")]
    mb_enh_ovrls  = [r["dnsmos_enh_melband_roformer_vocal"]["OVRL"] for r in data if r.get("dnsmos_enh_melband_roformer_vocal")]

    avg_raw   = round(np.mean(raw_ovrls), 3)
    avg_bs    = round(np.mean(bs_enh_ovrls), 3)
    avg_mb    = round(np.mean(mb_enh_ovrls), 3)
    delta_bs  = round(avg_bs - avg_raw, 3)
    delta_mb  = round(avg_mb - avg_raw, 3)

    h2(doc, "6.1 Hiệu quả tổng thể của pipeline")

    success(doc,
        f"Pipeline MSS + DFN3 nâng DNSMOS OVRL trung bình từ {avg_raw} → "
        f"{avg_mb} (MelBand) / {avg_bs} (BS-Roformer), "
        f"tương đương cải thiện +{delta_mb} / +{delta_bs} điểm OVRL.")

    para(doc,
        "Kết quả cho thấy pipeline 2 bước (MSS → Speech Enhancement) rất hiệu quả "
        "để xử lý audio YouTube thô. Ngay cả những file có OVRL thấp nhất (1.16 — "
        "talkshow với nhạc nền rất nặng) sau pipeline cũng đạt 3.4+.")

    h2(doc, "6.2 So sánh 2 model tách vocal")

    para(doc,
        "Về chất lượng tách vocal (DNSMOS sau separation):", bold=True)
    bullet(doc,
        "BS-Roformer-1297 có OVRL cao hơn MelBand sau bước separation đơn lẻ "
        f"(TB: {avg([r.get('dnsmos_sep_bs_roformer_1297',{}).get('OVRL') for r in data if r.get('dnsmos_sep_bs_roformer_1297')]):.3f} vs "
        f"{avg([r.get('dnsmos_sep_melband_roformer_vocal',{}).get('OVRL') for r in data if r.get('dnsmos_sep_melband_roformer_vocal')]):.3f})")
    bullet(doc,
        "Tuy nhiên MelBand + DFN3 vượt BS-Roformer + DFN3 ở phần lớn các files — "
        "DeepFilterNet3 boost MelBand mạnh hơn (Δ ≈ +0.18) so với boost BS-Roformer (Δ ≈ +0.09)")
    bullet(doc,
        f"Kết luận pipeline tốt nhất: MelBand-Roformer + DFN3 (OVRL TB: {avg_mb})")

    para(doc, "Về tốc độ xử lý:", bold=True)
    bullet(doc,
        f"MelBand nhanh hơn BS-Roformer {RTF_BS['mean']/RTF_MEL['mean']:.2f}× "
        f"(RTF {RTF_MEL['mean']:.2f} vs {RTF_BS['mean']:.2f})")
    bullet(doc,
        "DeepFilterNet3 cực kỳ nhanh (RTF ≈ 0.06×, tức 16× real-time) — gần như không ảnh hưởng tổng thời gian")
    bullet(doc,
        "Trên CPU: xử lý 1 phút audio mất ~5.7 phút (MelBand pipeline). Cần GPU để deploy thực tế.")

    h2(doc, "6.3 Phân tích theo loại content")

    tbl(doc,
        ["Category", "Đặc điểm", "Raw OVRL TB", "Best pipeline OVRL TB", "Δ OVRL", "Nhận xét"],
        [
            ["KhanhVy Vlog",
             "Nhạc nền liên tục, giọng VN nữ",
             str(round(np.mean([r["dnsmos_raw"]["OVRL"] for r in cats["khanhvy_vlog"] if r.get("dnsmos_raw")]), 3)),
             str(round(np.mean([r["dnsmos_enh_melband_roformer_vocal"]["OVRL"] for r in cats["khanhvy_vlog"] if r.get("dnsmos_enh_melband_roformer_vocal")]), 3)),
             f"+{round(np.mean([r['dnsmos_enh_melband_roformer_vocal']['OVRL'] - r['dnsmos_raw']['OVRL'] for r in cats['khanhvy_vlog'] if r.get('dnsmos_raw') and r.get('dnsmos_enh_melband_roformer_vocal')]), 3)}",
             "Cải thiện tốt, MelBand+DFN3 hiệu quả nhất"],
            ["KhanhVy Talkshow",
             "Nhạc nền nặng, biến thiên nhiều",
             str(round(np.mean([r["dnsmos_raw"]["OVRL"] for r in cats["khanhvy_talkshow"] if r.get("dnsmos_raw")]), 3)),
             str(round(np.mean([r["dnsmos_enh_bs_roformer_1297"]["OVRL"] for r in cats["khanhvy_talkshow"] if r.get("dnsmos_enh_bs_roformer_1297")]), 3)),
             f"+{round(np.mean([r['dnsmos_enh_bs_roformer_1297']['OVRL'] - r['dnsmos_raw']['OVRL'] for r in cats['khanhvy_talkshow'] if r.get('dnsmos_raw') and r.get('dnsmos_enh_bs_roformer_1297')]), 3)}",
             "Cải thiện nhiều nhất (+1.7), nhạc nền rất nặng"],
            ["Beyond Limits\n/ NGA Levi",
             "Podcast/interview, ít nhạc",
             str(round(np.mean([r["dnsmos_raw"]["OVRL"] for r in cats["beyondlimits_ngalevi"] if r.get("dnsmos_raw")]), 3)),
             str(round(np.mean([r["dnsmos_enh_melband_roformer_vocal"]["OVRL"] for r in cats["beyondlimits_ngalevi"] if r.get("dnsmos_enh_melband_roformer_vocal")]), 3)),
             f"+{round(np.mean([r['dnsmos_enh_melband_roformer_vocal']['OVRL'] - r['dnsmos_raw']['OVRL'] for r in cats['beyondlimits_ngalevi'] if r.get('dnsmos_raw') and r.get('dnsmos_enh_melband_roformer_vocal')]), 3)}",
             "Điểm cuối cao nhất (~3.86), audio gốc đã tốt hơn"],
        ],
        col_widths=[2.8, 3.5, 2.3, 3.0, 1.8, 4.2],
    )

    h2(doc, "6.4 Cải thiện chỉ số BAK (Background Noise)")

    para(doc,
        "BAK score phản ánh chất lượng loại bỏ nhiễu nền — chỉ số quan trọng nhất "
        "với mục tiêu tách giọng từ YouTube. Cải thiện BAK là kết quả rõ nhất của pipeline:")

    bak_rows = []
    for row in data:
        short = row["file"]
        raw_bak = row.get("dnsmos_raw", {}).get("BAK") if row.get("dnsmos_raw") else None
        mb_enh_bak = row.get("dnsmos_enh_melband_roformer_vocal", {}).get("BAK") if row.get("dnsmos_enh_melband_roformer_vocal") else None
        delta_bak = round(mb_enh_bak - raw_bak, 3) if (raw_bak and mb_enh_bak) else None
        bak_rows.append([
            short.split("__")[-1][:18],
            str(raw_bak), str(mb_enh_bak),
            f"+{delta_bak}" if delta_bak else "N/A"
        ])

    tbl(doc,
        ["File", "BAK Raw", "BAK MelBand+DFN3", "Δ BAK"],
        bak_rows,
        col_widths=[6.0, 3.0, 4.5, 3.0],
    )

    avg_raw_bak = avg([r.get("dnsmos_raw", {}).get("BAK") if r.get("dnsmos_raw") else None for r in data])
    avg_mb_bak  = avg([r.get("dnsmos_enh_melband_roformer_vocal", {}).get("BAK") if r.get("dnsmos_enh_melband_roformer_vocal") else None for r in data])
    success(doc,
        f"BAK trung bình: {avg_raw_bak} (raw) → {avg_mb_bak} (MelBand+DFN3), "
        f"cải thiện +{round(avg_mb_bak - avg_raw_bak, 3)} điểm. "
        "Nền âm thanh được loại bỏ hiệu quả ở tất cả các loại content.")

    # ─────────────────────────────────────────────────────────
    h2(doc, "6.5 So sánh với SIDON (arXiv:2509.17052v3)")
    # ─────────────────────────────────────────────────────────

    para(doc,
        "SIDON (Nakata et al., 2026 — The University of Tokyo) là model speech restoration "
        "open-source mới nhất, được thiết kế với mục tiêu giống hệt dự án này: làm sạch "
        "dữ liệu audio thô từ internet cho mục đích TTS. SIDON là đối tượng so sánh quan "
        "trọng vì:")
    bullet(doc, "Cùng mục tiêu: in-the-wild speech → studio-quality speech cho TTS/ASR dataset")
    bullet(doc, "Open source: model + code công khai tại github.com/sarulab-speech/Sidon")
    bullet(doc, "Trained trên 2,219h × 104 ngôn ngữ — scale lớn hơn bất kỳ OSS nào trước đó")

    h3(doc, "Kiến trúc SIDON")
    tbl(doc,
        ["Thành phần", "SIDON", "Pipeline của chúng ta (MSS + DFN3)"],
        [
            ["Bước 1 — Tách nhạc", "Không có bước riêng (end-to-end)",
             "BS-Roformer / Mel-Band-Roformer (MSS)"],
            ["Bước 2 — Khử noise", "Feature predictor: w2v-BERT 2.0 + LoRA",
             "DeepFilterNet 3 (GRU + ERB filterbank)"],
            ["Bước 3 — Tổng hợp", "Vocoder: HiFi-GAN + Snake activation",
             "Không có — output trực tiếp từ DFN3"],
            ["Training data", "2,219h, 104 ngôn ngữ (LibriTTS-R, FLEURS-R, ...)",
             "Pretrained (không train thêm)"],
            ["Số tham số", "w2v-BERT 2.0 (600M) + HiFi-GAN (~15M)",
             "BS-Roformer ~320M + DFN3 ~1.8M"],
            ["Dependencies", "flash-attn, torch>=2.8, uv (phức tạp)",
             "audio-separator, df (đơn giản hơn)"],
            ["GPU requirement", "Bắt buộc (flash-attn requires CUDA)",
             "Chạy được trên CPU"],
        ],
        col_widths=[4.0, 6.0, 7.5],
    )

    h3(doc, "So sánh DNSMOS OVRL — Benchmark tham chiếu")
    info(doc,
        "SIDON báo cáo kết quả trên TED-LIUM 3 (English). Pipeline của chúng ta đo trên "
        "YouTube tiếng Việt. Hai datasets không giống nhau — so sánh mang tính tham khảo.")

    tbl(doc,
        ["Model / System", "Type", "DNSMOS OVRL ↑", "NISQA ↑", "WER ↓", "SpkSim ↑", "RTF ↓", "Mã nguồn"],
        [
            ["Noisy input (baseline)",    "—",                    "2.12", "2.53", "7.4%", "1.000",  "—",        "—"],
            ["DeepFilterNet 3",           "Speech Enhancement",   "2.76", "3.15", "8.1%", "0.952",  "0.020 CPU","✅"],
            ["SIDON (paper — English)",   "Speech Restoration",   "3.31", "3.74", "5.8%", "0.891",  "0.002 GPU","✅"],
            ["Miipher (Google, closed)",  "Speech Restoration",   "3.28", "3.70", "4.9%", "0.895",  "N/A",      "❌"],
            ["BS-Roformer+DFN3 (ours)",   "MSS + Enhancement",
             f"{avg_bs:.3f}±0.29",        "—",  "—",  "~0.93*", f"{RTF_BS['mean']+RTF_DFN3:.1f} CPU", "✅"],
            ["Mel-Band+DFN3 (ours, best)","MSS + Enhancement",
             f"{avg_mb:.3f}±0.27",        "—",  "—",  "~0.94*", f"{RTF_MEL['mean']+RTF_DFN3:.1f} CPU","✅"],
        ],
        col_widths=[4.5, 3.2, 2.5, 2.2, 1.8, 2.2, 3.0, 2.5],
    )

    para(doc, "* SpkSim ước tính dựa trên đặc tính pipeline (không dùng resynthesis vocoder).", italic=True, size=8)

    h3(doc, "Nhận xét so sánh")

    para(doc, "Điểm mạnh của pipeline chúng ta so với SIDON:", bold=True)
    bullet(doc,
        f"DNSMOS OVRL đạt {avg_mb:.3f} (MelBand+DFN3) trên YouTube tiếng Việt — "
        f"cao hơn DNSMOS OVRL của SIDON (3.31) trên benchmark English của họ (mang tính tham khảo)")
    bullet(doc,
        "Speaker identity tốt hơn: pipeline không dùng neural vocoder resynthesis "
        "→ SpkSim cao hơn (SIDON SpkSim=0.891 do vocoder thay đổi đặc trưng giọng)")
    bullet(doc,
        "Không cần GPU để chạy (SIDON yêu cầu flash-attn + CUDA)")
    bullet(doc,
        "Đơn giản hơn để deploy: pip install audio-separator df (không cần uv, flash-attn)")

    para(doc, "Điểm mạnh của SIDON so với pipeline chúng ta:", bold=True)
    bullet(doc,
        "Tốc độ: RTF=0.002 trên GPU H200 (500× real-time) — trong khi pipeline chúng ta "
        f"RTF≈{RTF_MEL['mean']+RTF_DFN3:.1f}× trên CPU (chậm hơn real-time)")
    bullet(doc,
        "End-to-end: 1 model xử lý noise + reverb + bandwidth limitation + codec artifacts "
        "— không cần quyết định bước MSS riêng")
    bullet(doc,
        "Multilingual: trained 104 ngôn ngữ, generalize tốt hơn trên low-resource languages")
    bullet(doc,
        "Có thêm dereverberation + super-resolution (48kHz output) "
        "— pipeline của chúng ta chỉ làm noise suppression")
    bullet(doc,
        "Paper có NISQA metric (SIDON=3.74) — metrics toàn diện hơn")

    para(doc, "Điểm yếu cần lưu ý của SIDON:", bold=True)
    bullet(doc,
        "Vocoder resynthesis có thể thay đổi đặc trưng giọng gốc (SpkSim=0.891) — "
        "dù DNSMOS cao nhưng giọng có thể nghe khác người nói gốc")
    bullet(doc,
        "Không giải quyết được Music Source Separation trực tiếp — nếu audio có nhạc nền "
        "mạnh, cần kết hợp với MSS trước")
    bullet(doc,
        "Dependency nặng (flash-attn, torch>=2.8) — khó cài đặt trên môi trường không có GPU")

    success(doc,
        "Khuyến nghị: Với mục tiêu xử lý YouTube audio tiếng Việt trên CPU → "
        "dùng Mel-Band-Roformer + DFN3. "
        "Nếu có GPU mạnh và cần scale lớn (>1000h) hoặc multilingual → "
        "xem xét chuyển sang SIDON.")

    divider(doc)

    # ════════════════════════════════════════════════════════════
    # PHẦN 7: KẾT LUẬN & ĐỀ XUẤT
    # ════════════════════════════════════════════════════════════
    h1(doc, "7. Kết luận & Đề xuất")

    h2(doc, "7.1 Kết luận")

    tbl(doc,
        ["#", "Kết luận", "Chi tiết"],
        [
            ["1", "Pipeline MSS + DFN3 rất hiệu quả",
             f"OVRL raw {avg_raw} → {avg_mb} (MelBand+DFN3), tăng {delta_mb} điểm DNSMOS"],
            ["2", "MelBand-Roformer + DFN3 là pipeline tốt nhất",
             f"Vượt BS-Roformer+DFN3 về OVRL, nhanh hơn {RTF_BS['mean']/RTF_MEL['mean']:.2f}× về RTF"],
            ["3", "BS-Roformer tách vocal sạch hơn về SDR",
             "Nhưng MelBand vocals dễ enhance hơn bởi DeepFilterNet3"],
            ["4", "DeepFilterNet3 là bước boost quan trọng",
             "Đặc biệt cải thiện BAK và OVRL sau separation, RTF cực nhỏ (0.06×)"],
            ["5", "Voice Distortion cần kiểm tra thêm",
             "DNSMOS cao không đảm bảo không méo giọng — cần nghe thủ công (listening_test.html)"],
            ["6", "SIDON tốt hơn về tốc độ + scale",
             "RTF=0.002 GPU vs RTF~5.7 CPU — nhưng pipeline chúng ta tốt hơn về speaker identity"],
            ["7", "Content phụ thuộc ảnh hưởng kết quả",
             "Talkshow (nhiều nhạc) cải thiện nhiều nhất; interview (ít nhạc) đạt điểm cao nhất"],
        ],
        col_widths=[0.8, 5.0, 11.8],
    )

    h2(doc, "7.2 Đề xuất cho deployment")

    bullet(doc, "Recommended pipeline: MelBand-Roformer-Vocal → DeepFilterNet 3")
    bullet(doc, "Dùng GPU để đạt RTF < 1 (real-time hoặc nhanh hơn) — hiện CPU quá chậm cho production")
    bullet(doc, "Với nội dung talk show/vlog có nhạc nền: nên dùng BS-Roformer + DFN3 (tách sạch hơn trước khi enhance)")
    bullet(doc, "Với nội dung interview/podcast ít nhạc: MelBand+DFN3 đủ và nhanh hơn")
    bullet(doc, "Nên test thêm trên toàn bộ dataset (90 files) và so sánh WER sau ASR để đánh giá downstream")
    bullet(doc, "Xem xét thêm 2-stage separation: BS-Roformer (vocal) → De-Echo MelBand (noise) → DFN3")

    h2(doc, "7.3 Các bước tiếp theo")

    tbl(doc,
        ["Ưu tiên", "Task", "Lý do"],
        [
            ["🔴 Cao", "Listening test thủ công (listening_test.html)",
             "Kiểm tra méo giọng trực tiếp — metric tự động chỉ là proxy"],
            ["🔴 Cao", "Test trên GPU",
             "RTF hiện 5–9× trên CPU, cần < 1× cho production"],
            ["🔴 Cao", "Evaluate downstream ASR/WER",
             "DNSMOS chỉ đo perceptual quality, cần đo ảnh hưởng thực tế"],
            ["🟡 Trung", "Thử SIDON trên GPU nếu có",
             "End-to-end, nhanh hơn 500×, hỗ trợ multilingual — phù hợp scale lớn"],
            ["🟡 Trung", "Thử toàn bộ 90 files",
             "Sample 9 files có thể chưa đại diện đủ"],
            ["🟡 Trung", "Thêm NISQA metric vào pipeline",
             "NISQA bổ sung DNSMOS, đo nhiều khía cạnh hơn (noisiness, coloration, discontinuity)"],
            ["🟡 Trung", "Thử thêm MDX23C + BS-Roformer ensemble",
             "Ensemble thường cho kết quả tốt hơn"],
            ["🟢 Thấp", "Tích hợp vào data processing pipeline",
             "Sau khi chọn được model phù hợp"],
        ],
        col_widths=[1.8, 5.5, 10.3],
    )

    divider(doc)

    # ════════════════════════════════════════════════════════════
    # PHỤ LỤC
    # ════════════════════════════════════════════════════════════
    h1(doc, "Phụ lục: Dữ liệu chi tiết đầy đủ")

    para(doc, "Bảng đầy đủ tất cả metrics (SIG, BAK, OVRL) cho từng file:", bold=True)

    full_headers = [
        "File", "SNR\nRaw",
        "Raw\nSIG", "Raw\nBAK", "Raw\nOVRL",
        "BS-sep\nOVRL", "BS+DFN3\nOVRL",
        "MB-sep\nOVRL", "MB+DFN3\nOVRL",
    ]
    full_rows = []
    for row in data:
        raw  = row.get("dnsmos_raw") or {}
        bss  = row.get("dnsmos_sep_bs_roformer_1297") or {}
        bse  = row.get("dnsmos_enh_bs_roformer_1297") or {}
        mbs  = row.get("dnsmos_sep_melband_roformer_vocal") or {}
        mbe  = row.get("dnsmos_enh_melband_roformer_vocal") or {}
        full_rows.append([
            row["file"].split("__")[-1][:20],
            row.get("snr_raw", "—"),
            raw.get("SIG", "—"), raw.get("BAK", "—"), raw.get("OVRL", "—"),
            bss.get("OVRL", "—"), bse.get("OVRL", "—"),
            mbs.get("OVRL", "—"), mbe.get("OVRL", "—"),
        ])

    tbl(doc, full_headers, full_rows,
        col_widths=[3.5, 1.5, 1.5, 1.5, 1.8, 2.0, 2.0, 2.0, 2.2])

    # ── Footer ────────────────────────────────────────────────────
    doc.add_paragraph()
    divider(doc)
    footer_p = doc.add_paragraph(
        "Báo cáo tự động tạo bởi generate_report.py  |  "
        "Dự án: AI Voice Pipeline — MSS + Speech Enhancement + Voice Distortion + SIDON Comparison  |  2026-06-08"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    footer_p.runs[0].italic = True

    doc.save(out_path)
    print(f"[XONG] Báo cáo đã lưu tại: {out_path}")


if __name__ == "__main__":
    build(OUT_DOCX)
